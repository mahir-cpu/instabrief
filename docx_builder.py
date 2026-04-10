import os
import tempfile
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2E, 0x86, 0xC1)
CHARCOAL = RGBColor(0x2D, 0x2D, 0x2D)
BODY_COLOR = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ROW_BG = "E8EDF2"

# Total usable width with 2cm margins on A4 = ~6.69 inches
TOTAL_WIDTH = 6.69


def _label_width(labels):
    longest = max(len(label) for label in labels) if labels else 10
    width = longest * 0.09 + 0.3
    return min(max(width, 1.0), 2.5)


def _set_font(run, name="Calibri", size=Pt(10.5), color=BODY_COLOR, bold=False):
    run.font.name = name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold


def _add_bottom_border(paragraph, color="2E86C1", width="4"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), width)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_hyperlink(paragraph, text, url, color=BODY_COLOR, size=Pt(10.5), bold=False):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "2E86C1")
    rPr.append(c)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size.pt * 2)))
    rPr.append(sz)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell, color="D0D5DD"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ["top", "left", "bottom", "right"]:
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_width(cell, inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_cell_valign_top(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "top")
    tcPr.append(vAlign)


def _cell_add_text(cell, text, size=Pt(9), color=BODY_COLOR, bold=False):
    p = cell.paragraphs[0] if cell.paragraphs and cell.paragraphs[0].text == "" else cell.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=size, color=color, bold=bold)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    return p


def _cell_add_bullet(cell, text, size=Pt(9), color=BODY_COLOR):
    p = cell.add_paragraph()
    run = p.add_run("\u2022 " + text)
    _set_font(run, size=size, color=color)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.15
    return p


def _add_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=Pt(16), color=NAVY, bold=True)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    _add_bottom_border(p)
    return p


def _add_body(doc, text, justify=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _add_body_italic(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, color=GRAY)
    run.italic = True
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def _add_labeled_line(doc, label, text):
    p = doc.add_paragraph()
    label_run = p.add_run(label + " ")
    _set_font(label_run, bold=True)
    text_run = p.add_run(text)
    _set_font(text_run)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    return p


def _add_two_col_table(doc, rows_data):
    labels = [label for label, _ in rows_data]
    left_w = _label_width(labels)
    right_w = TOTAL_WIDTH - left_w

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, (label, value) in enumerate(rows_data):
        row = table.add_row()
        cells = row.cells
        bg = ROW_BG if idx % 2 == 0 else "FFFFFF"

        _set_cell_width(cells[0], left_w)
        _set_cell_shading(cells[0], bg)
        _set_cell_borders(cells[0])
        _set_cell_valign_top(cells[0])
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(label)
        _set_font(run, size=Pt(9.5), color=NAVY, bold=True)

        _set_cell_width(cells[1], right_w)
        _set_cell_shading(cells[1], bg)
        _set_cell_borders(cells[1])
        _set_cell_valign_top(cells[1])
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(value)
        _set_font(run, size=Pt(9.5))

    return table


def _add_attendee_table(doc, person):
    name = person.get("name", "")
    linkedin = person.get("linkedin_url", "")
    position = person.get("current_position", "")

    rows_data = [("Name", "")]
    if person.get("career_history"):
        rows_data.append(("Career History", person["career_history"]))
    if person.get("education"):
        rows_data.append(("Education", person["education"]))
    if person.get("background"):
        rows_data.append(("Background", person["background"]))
    if person.get("past_call_context"):
        rows_data.append(("From past calls", person["past_call_context"]))

    labels = [label for label, _ in rows_data]
    left_w = _label_width(labels)
    right_w = TOTAL_WIDTH - left_w

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, (label, value) in enumerate(rows_data):
        row = table.add_row()
        cells = row.cells
        bg = ROW_BG if idx % 2 == 0 else "FFFFFF"

        _set_cell_width(cells[0], left_w)
        _set_cell_shading(cells[0], bg)
        _set_cell_borders(cells[0])
        _set_cell_valign_top(cells[0])
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(label)
        _set_font(run, size=Pt(9.5), color=NAVY, bold=True)

        _set_cell_width(cells[1], right_w)
        _set_cell_shading(cells[1], bg)
        _set_cell_borders(cells[1])
        _set_cell_valign_top(cells[1])
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15

        if label == "Name":
            if linkedin:
                _add_hyperlink(p, name, linkedin, NAVY, Pt(9.5), bold=True)
            else:
                run = p.add_run(name)
                _set_font(run, size=Pt(9.5), color=NAVY, bold=True)
            if position:
                run = p.add_run(" -- " + position)
                _set_font(run, size=Pt(9.5), color=GRAY)
        else:
            run = p.add_run(value)
            _set_font(run, size=Pt(9.5))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)

    return table


def _add_relationship_table(doc, rel_history):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Meeting", "Key highlights", "Outcome", "Next steps"]
    widths = [1.5, 2.2, 1.7, 1.6]

    header_cells = table.rows[0].cells
    for i, (cell, header, width) in enumerate(zip(header_cells, headers, widths)):
        _set_cell_width(cell, width)
        _set_cell_shading(cell, "1B2A4A")
        _set_cell_borders(cell, "1B2A4A")
        _set_cell_valign_top(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(header)
        _set_font(run, size=Pt(9), color=WHITE, bold=True)

    for idx, meeting in enumerate(rel_history):
        row = table.add_row()
        cells = row.cells
        bg = ROW_BG if idx % 2 == 0 else "FFFFFF"

        for i, cell in enumerate(cells):
            _set_cell_width(cell, widths[i])
            _set_cell_shading(cell, bg)
            _set_cell_borders(cell)
            _set_cell_valign_top(cell)

        label = meeting.get("meeting_label", "")
        meeting_date = meeting.get("meeting_date", "")
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(label)
        _set_font(run, size=Pt(9), color=NAVY, bold=True)
        if meeting_date:
            p2 = cells[0].add_paragraph()
            run2 = p2.add_run(meeting_date)
            _set_font(run2, size=Pt(8), color=GRAY)
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(2)

        highlights = meeting.get("key_highlights", [])
        if highlights:
            p = cells[1].paragraphs[0]
            run = p.add_run("\u2022 " + highlights[0])
            _set_font(run, size=Pt(8.5), color=BODY_COLOR)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.15
            for h in highlights[1:]:
                _cell_add_bullet(cells[1], h, size=Pt(8.5), color=BODY_COLOR)
        else:
            _cell_add_text(cells[1], "N/A", size=Pt(8.5))

        outcome = meeting.get("outcome", "")
        _cell_add_text(cells[2], outcome, size=Pt(8.5))

        next_step = meeting.get("next_step", "")
        _cell_add_text(cells[3], next_step, size=Pt(8.5))


def build_docx(data):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_date.add_run("Prepared for InstaBrief  |  " + date.today().strftime("%B %d, %Y"))
    _set_font(run, size=Pt(10), color=GRAY)
    run.italic = True

    p_title = doc.add_paragraph()
    run = p_title.add_run(data["company_name"])
    _set_font(run, size=Pt(28), color=NAVY, bold=True)

    p_div = doc.add_paragraph()
    _add_bottom_border(p_div, color="2E86C1", width="6")
    p_div.paragraph_format.space_after = Pt(6)

    ctx = data.get("company_context", "")
    if ctx:
        _add_body_italic(doc, ctx)

    # Meeting Attendees
    attendees = data.get("meeting_attendees", [])
    if attendees:
        _add_section_header(doc, "Meeting Attendees")
        for person in attendees:
            _add_attendee_table(doc, person)

    # Relationship History as table
    rel_history = data.get("relationship_history", [])
    if rel_history:
        _add_section_header(doc, "Relationship History")
        _add_relationship_table(doc, rel_history)

    # Next Steps and Objections
    next_steps = data.get("next_steps", "")
    objections = data.get("objections", "")
    if next_steps or objections:
        _add_section_header(doc, "Next Steps & Objections")
        if next_steps:
            _add_labeled_line(doc, "Next Steps:", next_steps)
        if objections:
            _add_labeled_line(doc, "Key Objections:", objections)

    # Client Profile as table
    _add_section_header(doc, "Client Profile")
    profile = data.get("client_profile", {})
    profile_rows = [
        ("What they do", profile.get("what_they_do", "N/A")),
        ("Markets served", profile.get("markets_served", "N/A")),
        ("Revenue", profile.get("revenue", "N/A")),
        ("Scale", profile.get("scale", "N/A")),
        ("Recent growth", profile.get("recent_growth", "N/A")),
    ]
    _add_two_col_table(doc, profile_rows)

    # Core Pain Points as table
    _add_section_header(doc, "Core Pain Points")
    pain_rows = []
    for pp in data.get("core_pain_points", []):
        pain_rows.append((pp.get("title", ""), pp.get("description", "")))
    if pain_rows:
        _add_two_col_table(doc, pain_rows)

    # Highest-Impact Solutions as table
    _add_section_header(doc, "Highest-Impact Agentic AI Solutions")
    sol_rows = []
    for sol in data.get("highest_impact_solutions", []):
        sol_rows.append((sol.get("name", ""), sol.get("description", "")))
    if sol_rows:
        _add_two_col_table(doc, sol_rows)

    # Best Approach
    _add_section_header(doc, "Best Approach")
    ba = data.get("best_approach", "")
    if isinstance(ba, list):
        for para in ba:
            _add_body(doc, para)
    else:
        _add_body(doc, ba)

    # AI Insight
    _add_section_header(doc, "Relevant AI-Related Insight")
    ai = data.get("ai_insight", "")
    if isinstance(ai, list):
        for para in ai:
            _add_body(doc, para)
    else:
        _add_body(doc, ai)

    # Save
    safe_name = data.get("company_name", "Brief").replace(" ", "_").replace("/", "-")
    filepath = os.path.join(tempfile.gettempdir(), safe_name + "_InstaBrief.docx")
    doc.save(filepath)
    return filepath